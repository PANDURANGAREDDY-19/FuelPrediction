"""
Resume Parser - Extract and clean text from PDF, DOCX, and TXT files
"""
import re
from pathlib import Path
import PyPDF2
import pdfplumber
from docx import Document

class ResumeParser:
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt']
    
    def parse(self, file_path):
        """Parse resume and extract clean text"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        ext = file_path.suffix.lower()
        if ext not in self.supported_formats:
            raise ValueError(f"Unsupported format: {ext}. Supported: {self.supported_formats}")
        
        # Route to appropriate parser
        parsers = {'.pdf': self._parse_pdf, '.docx': self._parse_docx, '.txt': self._parse_txt}
        return parsers[ext](file_path)
    
    def _parse_pdf(self, file_path):
        """Extract text from PDF using pdfplumber (primary) and PyPDF2 (fallback)"""
        text = ""
        
        # Try pdfplumber first (better text extraction)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception:
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            except Exception:
                pass
        
        return self._clean_text(text)
    
    def _parse_docx(self, file_path):
        """Extract text from DOCX including paragraphs and tables"""
        doc = Document(file_path)
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        return self._clean_text('\n'.join(text_parts))
    
    def _parse_txt(self, file_path):
        """Extract text from TXT with encoding fallback"""
        encodings = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return self._clean_text(f.read())
            except UnicodeDecodeError:
                continue
        
        raise ValueError(f"Unable to decode file with encodings: {encodings}")
    
    def _clean_text(self, text):
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Normalize line breaks
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove excessive whitespace while preserving structure
        text = re.sub(r' +', ' ', text)  # Multiple spaces to single
        text = re.sub(r'\n{3,}', '\n\n', text)  # Max 2 consecutive newlines
        
        # Remove special characters but keep important punctuation
        text = re.sub(r'[^\w\s@.,\-+()#:/\n]', '', text)
        
        return text.strip()
